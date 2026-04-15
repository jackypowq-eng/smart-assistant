import os
import logging
from typing import List, Optional
from langchain_core.documents import Document

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self, embedding_model=None):
        self.embedding_model = embedding_model
        self.vector_store = None
        try:
            self._init_embedding_model()
        except Exception as e:
            logger.warning(f"初始化 Embedding 模型失败: {str(e)}")
            logger.warning("向量化功能将不可用，但文档解析功能仍然可用")
        
    def _init_embedding_model(self):
        if self.embedding_model is None:
            self.embedding_model = self._get_embedding_model()
    
    def _get_embedding_model(self):
        try:
            from langchain_community.embeddings import HuggingFaceEmbeddings
            logger.info("使用本地 HuggingFace Embeddings")
            model = HuggingFaceEmbeddings(
                model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2",
                model_kwargs={'device': 'cpu'},
                encode_kwargs={'normalize_embeddings': True}
            )
            logger.info("HuggingFace Embeddings 初始化完成")
            return model
        except Exception as e:
            logger.error(f"初始化本地 embedding 模型失败: {str(e)}", exc_info=True)
            raise
        
    def _read_file_content(self, file_path: str) -> str:
        encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16', 'latin-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                    if content.strip():
                        return content
            except (UnicodeDecodeError, UnicodeError):
                continue
            except Exception as e:
                logger.warning(f"读取文件 {file_path} 时出错 (编码: {encoding}): {str(e)}")
                continue
        
        try:
            with open(file_path, 'rb') as f:
                content = f.read().decode('utf-8', errors='ignore')
                return content
        except Exception as e:
            raise Exception(f"无法读取文件 {file_path}: {str(e)}")
    
    def _extract_text_from_pdf(self, file_path: str) -> List[str]:
        try:
            from langchain_community.document_loaders import PyPDFLoader
            loader = PyPDFLoader(file_path)
            pages = loader.load()
            logger.info(f"PDF 加载完成: {file_path}, 共 {len(pages)} 页")
            return [page.page_content for page in pages if page.page_content.strip()]
        except Exception as e:
            raise Exception(f"PDF 解析失败: {str(e)}")
    
    def _extract_text_from_docx(self, file_path: str) -> List[str]:
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            logger.info(f"DOCX 加载完成: {file_path}, 共 {len(paragraphs)} 段")
            return paragraphs
        except ImportError:
            try:
                from langchain_community.document_loaders import UnstructuredWordDocumentLoader
                loader = UnstructuredWordDocumentLoader(file_path)
                docs = loader.load()
                return [doc.page_content for doc in docs if doc.page_content.strip()]
            except Exception as e:
                raise Exception(f"DOCX 解析失败 (请安装 python-docx): {str(e)}")
        except Exception as e:
            raise Exception(f"DOCX 解析失败: {str(e)}")
    
    def _extract_text_from_xlsx(self, file_path: str) -> List[str]:
        """从 Excel 提取文本"""
        try:
            import pandas as pd
            df = pd.read_excel(file_path, sheet_name=None)
            texts = []
            for sheet_name, sheet_df in df.items():
                for _, row in sheet_df.iterrows():
                    row_text = ' '.join([str(cell) for cell in row if pd.notna(cell)])
                    if row_text.strip():
                        texts.append(row_text)
            return texts
        except ImportError:
            raise Exception("Excel 解析失败: 请安装 pandas 和 openpyxl")
        except Exception as e:
            raise Exception(f"Excel 解析失败: {str(e)}")
    
    def _extract_text_from_csv(self, file_path: str) -> List[str]:
        """从 CSV 提取文本"""
        try:
            import pandas as pd
            df = pd.read_csv(file_path, encoding='utf-8')
        except UnicodeDecodeError:
            df = pd.read_csv(file_path, encoding='gbk')
        
        texts = []
        for _, row in df.iterrows():
            row_text = ' '.join([str(cell) for cell in row if pd.notna(cell)])
            if row_text.strip():
                texts.append(row_text)
        return texts
    
    def _split_text(self, text: str, chunk_size: int = 500, chunk_overlap: int = 50) -> List[str]:
        """智能文本分割"""
        if not text or not text.strip():
            return []
        
        # 按段落分割
        paragraphs = text.split('\n\n')
        chunks = []
        current_chunk = ""
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # 如果当前块加上新段落不超过限制，则添加
            if len(current_chunk) + len(para) + 2 <= chunk_size:
                current_chunk += "\n\n" + para if current_chunk else para
            else:
                # 保存当前块
                if current_chunk:
                    chunks.append(current_chunk)
                
                # 如果段落本身超过限制，需要进一步分割
                if len(para) > chunk_size:
                    # 按句子分割
                    sentences = []
                    for sep in ['。', '！', '？', '.', '!', '?', '\n']:
                        parts = para.split(sep)
                        if len(parts) > 1:
                            sentences = [p + sep for p in parts[:-1]] + [parts[-1]]
                            break
                    
                    if not sentences:
                        sentences = [para]
                    
                    temp_chunk = ""
                    for sentence in sentences:
                        if len(temp_chunk) + len(sentence) <= chunk_size:
                            temp_chunk += sentence
                        else:
                            if temp_chunk:
                                chunks.append(temp_chunk)
                            temp_chunk = sentence
                    
                    if temp_chunk:
                        current_chunk = temp_chunk
                    else:
                        current_chunk = ""
                else:
                    current_chunk = para
        
        # 添加最后一块
        if current_chunk:
            chunks.append(current_chunk)
        
        return chunks
    
    def process_file(self, file_path: str) -> List[Document]:
        """处理单个文件"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        # 检查文件大小
        file_size = os.path.getsize(file_path)
        if file_size == 0:
            raise ValueError(f"文件为空: {file_path}")
        
        # 获取文件扩展名
        ext = os.path.splitext(file_path)[1].lower()
        filename = os.path.basename(file_path)
        
        # 提取文本
        texts = []
        try:
            if ext == '.pdf':
                texts = self._extract_text_from_pdf(file_path)
            elif ext == '.docx':
                texts = self._extract_text_from_docx(file_path)
            elif ext == '.xlsx':
                texts = self._extract_text_from_xlsx(file_path)
            elif ext == '.csv':
                texts = self._extract_text_from_csv(file_path)
            elif ext in ['.txt', '.md']:
                content = self._read_file_content(file_path)
                texts = [content] if content.strip() else []
            else:
                # 尝试作为文本文件读取
                content = self._read_file_content(file_path)
                texts = [content] if content.strip() else []
        except Exception as e:
            raise Exception(f"文件解析失败 ({filename}): {str(e)}")
        
        if not texts:
            raise ValueError(f"文件内容为空或无法解析: {filename}")
        
        # 分块处理
        documents = []
        for text in texts:
            if not text.strip():
                continue
            
            chunks = self._split_text(text)
            for i, chunk in enumerate(chunks):
                if chunk.strip():
                    doc = Document(
                        page_content=chunk,
                        metadata={
                            'source': filename,
                            'file_path': file_path,
                            'chunk_index': i
                        }
                    )
                    documents.append(doc)
        
        return documents
    
    def load_and_process(self, file_paths: List[str]) -> int:
        all_documents = []
        errors = []
        
        for file_path in file_paths:
            try:
                docs = self.process_file(file_path)
                all_documents.extend(docs)
                logger.info(f"处理完成: {os.path.basename(file_path)} ({len(docs)} 个片段)")
            except Exception as e:
                error_msg = f"处理失败: {os.path.basename(file_path)} - {str(e)}"
                logger.error(error_msg)
                errors.append(error_msg)
        
        if not all_documents:
            if errors:
                raise Exception("所有文件处理失败:\n" + "\n".join(errors))
            else:
                raise Exception("没有有效的文档内容")
        
        try:
            from langchain_community.vectorstores import Chroma
            
            logger.info(f"开始向量化，共 {len(all_documents)} 个文档片段...")
            self.vector_store = Chroma.from_documents(
                documents=all_documents,
                embedding=self.embedding_model,
                persist_directory="./chroma_db"
            )
            
            logger.info(f"向量化完成: 共 {len(all_documents)} 个文档片段")
            
        except Exception as e:
            logger.error(f"向量化失败，但文档已解析: {str(e)}")
            logger.warning("文档解析成功，向量化失败 - 文档状态仍将标记为已处理")
        
        return len(all_documents)
    
    def retrieve(self, query: str, k: int = 4) -> List[Document]:
        if not self.vector_store:
            return []
        
        try:
            return self.vector_store.similarity_search(query, k=k)
        except Exception as e:
            logger.error(f"检索失败: {str(e)}")
            return []
